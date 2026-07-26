# app/services/document_gen.py

from docx import Document
import hashlib, io, os

TEMPLATES = {
    'chargesheet':       'templates/documents/chargesheet_bns2024.docx',
    'medical_letter':    'templates/documents/medical_letter.docx',
    'remand_request':    'templates/documents/remand_request_bnss.docx',
    'seizure_receipt':   'templates/documents/seizure_receipt.docx',
    'court_custody':     'templates/documents/court_custody_bnss.docx',
    'panchanama':        'templates/documents/accused_panchanama.docx',
    'face_id':           'templates/documents/face_identification.docx',
    'witness_statement': 'templates/documents/witness_statement.docx',
    'fir':               'templates/documents/chargesheet_bns2024.docx',
    'case_diary':        'templates/documents/accused_panchanama.docx',
    'arrest_memo':       'templates/documents/accused_panchanama.docx',
    'seizure_list':      'templates/documents/seizure_receipt.docx',
    'search_warrant':    'templates/documents/court_custody_bnss.docx',
    'bail_objection':    'templates/documents/remand_request_bnss.docx',
}

# GUJARATI DOMAIN GLOSSARY — override before IndicTrans2
GLOSSARY = {
    'FIR':          {'gu': 'ફ.ઇ.ર',       'hi': 'प्रथम सूचना रिपोर्ट'},
    'chargesheet':  {'gu': 'ચાજ્જશીટ',     'hi': 'आरोप पत्र'},
    'accused':      {'gu': 'આરોપી',        'hi': 'अभियुक्त'},
    'seizure':      {'gu': 'જપ્તી',         'hi': 'जब्ती'},
    'remand':       {'gu': 'રિમાન્ડ',       'hi': 'रिमांड'},
    'panchanama':   {'gu': 'પંચનામું',      'hi': 'पंचनामा'},
    'witness':      {'gu': 'એાક્ષી',        'hi': 'गवाह'},
    'arrest':       {'gu': 'ધરપકડ',        'hi': 'गिरफ्तारी'},
    'police station':{'gu': 'પોલીએ સ્ટેશન','hi': 'थाना'},
}

def replace_in_paragraph(para, ctx):
    for key, val in ctx.items():
        tag = '{{' + key + '}}'
        val_str = str(val) if val is not None else ''

        while True:
            if not para.runs:
                if tag in para.text:
                    para.text = para.text.replace(tag, val_str)
                break

            full_text = "".join(r.text for r in para.runs)
            tag_start = full_text.find(tag)
            if tag_start == -1:
                break

            tag_end = tag_start + len(tag)

            # Build run character offset mapping
            run_spans = []
            char_cursor = 0
            for run in para.runs:
                r_len = len(run.text)
                run_spans.append((run, char_cursor, char_cursor + r_len))
                char_cursor += r_len

            # Find start and end run indices
            start_run_idx = None
            end_run_idx = None
            for idx, (run, rstart, rend) in enumerate(run_spans):
                if start_run_idx is None and rstart <= tag_start < rend:
                    start_run_idx = idx
                if rend >= tag_end and rstart < tag_end:
                    end_run_idx = idx
                    break

            if start_run_idx is not None and end_run_idx is not None:
                if start_run_idx == end_run_idx:
                    run, rstart, _ = run_spans[start_run_idx]
                    l_start = tag_start - rstart
                    l_end = tag_end - rstart
                    run.text = run.text[:l_start] + val_str + run.text[l_end:]
                else:
                    # First run
                    first_run, rstart_first, _ = run_spans[start_run_idx]
                    l_start = tag_start - rstart_first
                    first_run.text = first_run.text[:l_start] + val_str

                    # Middle runs
                    for mid_idx in range(start_run_idx + 1, end_run_idx):
                        run_spans[mid_idx][0].text = ""

                    # Last run
                    last_run, rstart_last, _ = run_spans[end_run_idx]
                    l_end = tag_end - rstart_last
                    last_run.text = last_run.text[l_end:]
            else:
                # Fallback if run span mapping failed unexpectedly
                if tag in para.text:
                    full_text = para.text.replace(tag, val_str)
                    para.runs[0].text = full_text
                    for run in para.runs[1:]:
                        run.text = ""
                break

def generate_document(
    doc_type: str, 
    case: dict, 
    officer: dict,
    lang: str = 'en'
) -> tuple[bytes, str]:
    
    sections = case.get('bns_sections', [])
    
    # Context dictionary — all placeholders
    ctx = {
        'FIR_NO':           case.get('fir_no', ''),
        'DATE':             today_formatted(),
        'PS_NAME':          translate(case.get('ps_name', ''), lang),
        'IO_NAME':          officer.get('name', ''),
        'IO_BADGE':         officer.get('badge_no', ''),
        'VICTIM_NAME':      translate(case.get('victim_name', ''), lang),
        'VICTIM_ADDRESS':   translate(case.get('victim_address', ''), lang),
        'VICTIM_AGE':       str(case.get('victim_age', '')),
        'VICTIM_PHONE':     case.get('victim_phone', ''),
        'ACCUSED_NAME':     translate(case.get('accused_name', ''), lang),
        'ACCUSED_ADDRESS':  translate(case.get('accused_address', ''), lang),
        'ACCUSED_AGE':      str(case.get('accused_age', '')),
        'CRIME_TYPE':       translate(case.get('crime_type', ''), lang),
        'CRIME_DATE':       format_date(case.get('crime_date')),
        'CRIME_LOCATION':   translate(case.get('crime_location', ''), lang),
        'CRIME_NARRATIVE':  translate(case.get('crime_narrative', ''), lang),
        'BNS_SECTIONS':     ', '.join(case.get('bns_sections') or []),
        'BNSS_SECTIONS':    ', '.join(case.get('bnss_sections') or []),
        'BSA_SECTIONS':     ', '.join(case.get('bsa_sections') or []),
        'IPC_CROSSREF':     ', '.join(case.get('ipc_crossref') or []),
        'LANDMARK_CASES':   format_landmark_cases(case.get('landmark_cases') or []),
        'EVIDENCE_LIST':    format_evidence(case.get('evidence_items') or [], lang),
        'WITNESS_LIST':     format_witnesses(case.get('witnesses') or [], lang),
        'ARREST_DATE':      format_date(case.get('arrest_date')),
        'ARREST_LOCATION':  case.get('arrest_location', ''),
    }
    
    template_path = TEMPLATES[doc_type]
    if not os.path.exists(template_path):
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        alt_path = os.path.join(base_dir, template_path)
        if os.path.exists(alt_path):
            template_path = alt_path

    doc = Document(template_path)
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, ctx)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, ctx)
    
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    sha256 = hashlib.sha256(raw).hexdigest()
    
    return raw, sha256

# CRITICAL: All templates use BNS/BNSS/BSA 2024 ONLY
# IPC/CrPC appear only in cross-reference column
# Ctrl+F "Section " in every template before demo
# Any "IPC" or "CrPC" as primary section = court rejection risk

from datetime import datetime, timezone

def today_formatted():
    return datetime.now(timezone.utc).strftime('%Y-%m-%d')

def format_date(d):
    return d.strftime('%Y-%m-%d') if d else ''

def translate(text, lang):
    from app.services.translation import translator_service
    if not text: return ''
    if lang == 'en': return text
    if hasattr(translator_service, 'translate_sync'):
        return translator_service.translate_sync(text, lang)
    return translator_service.translate(text, lang)

def format_landmark_cases(cases):
    if not cases: return ''
    return ', '.join([c if isinstance(c, str) else str(c) for c in cases])

def format_evidence(evidence, lang):
    if not evidence: return ''
    text = ', '.join([e if isinstance(e, str) else str(e) for e in evidence])
    return translate(text, lang)

def format_witnesses(witnesses, lang):
    if not witnesses: return ''
    text = ', '.join([w if isinstance(w, str) else str(w) for w in witnesses])
    return translate(text, lang)
