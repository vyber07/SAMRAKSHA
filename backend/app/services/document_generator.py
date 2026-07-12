import os
from docx import Document
from datetime import datetime

DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "generated_docs")
os.makedirs(DOCS_DIR, exist_ok=True)

def generate_purvani_chargesheet(case_data: dict) -> str:
    doc = Document()
    doc.add_heading('Purvani Chargesheet', 0)
    
    doc.add_paragraph(f"FIR Number: {case_data.get('fir_number')}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Title: {case_data.get('title')}")
    
    doc.add_heading('BNS Sections', level=1)
    for section in case_data.get('bns_sections', []):
        doc.add_paragraph(f"- {section}", style='List Bullet')
        
    doc.add_heading('Description', level=1)
    doc.add_paragraph(case_data.get('description', ''))
    
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_purvani_chargesheet.docx")
    doc.save(filepath)
    return filepath

def generate_medical_treatment_letter(case_data: dict) -> str:
    doc = Document()
    doc.add_heading('Medical Treatment Letter', 0)
    doc.add_paragraph(f"Reference FIR: {case_data.get('fir_number')}")
    doc.add_paragraph(f"Victim Injury Flag: {case_data.get('victim_injury')}")
    if case_data.get('victim_injury'):
        doc.add_paragraph("Urgent medical examination and treatment required.")
    
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_medical_letter.docx")
    doc.save(filepath)
    return filepath

def generate_seizure_receipt(case_data: dict, evidences: list) -> str:
    doc = Document()
    doc.add_heading('Seizure Receipt', 0)
    doc.add_paragraph(f"Case: {case_data.get('fir_number')}")
    
    doc.add_heading('Evidence Seized', level=1)
    for ev in evidences:
        doc.add_paragraph(f"- {ev.get('name')}: {ev.get('description')} (Hash: {ev.get('file_hash')})", style='List Bullet')
        
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_seizure_receipt.docx")
    doc.save(filepath)
    return filepath

def generate_remand_request(case_data: dict) -> str:
    doc = Document()
    doc.add_heading('Remand Request (Police Custody)', 0)
    doc.add_paragraph(f"FIR: {case_data.get('fir_number')}")
    doc.add_paragraph("Subject: Request for police custody remand of accused.")
    doc.add_paragraph("Under BNSS 2024 Sections...")
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_remand_request.docx")
    doc.save(filepath)
    return filepath

def generate_court_custody_letter(case_data: dict) -> str:
    doc = Document()
    doc.add_heading('Court Custody Letter', 0)
    doc.add_paragraph(f"FIR: {case_data.get('fir_number')}")
    doc.add_paragraph("As per BNSS 184 compliant template, requesting judicial custody.")
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_court_custody.docx")
    doc.save(filepath)
    return filepath

def generate_accused_panchanama(case_data: dict, witnesses: list) -> str:
    doc = Document()
    doc.add_heading('Accused Panchanama', 0)
    doc.add_paragraph(f"FIR: {case_data.get('fir_number')}")
    doc.add_heading('Witnesses', level=1)
    for w in witnesses:
        doc.add_paragraph(f"- {w}", style='List Bullet')
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_panchanama.docx")
    doc.save(filepath)
    return filepath

def generate_face_id_form(case_data: dict) -> str:
    doc = Document()
    doc.add_heading('Face Identification Form', 0)
    doc.add_paragraph(f"FIR: {case_data.get('fir_number')}")
    doc.add_paragraph("[ Photo Upload Box Placeholder ]")
    doc.add_paragraph("Description: ")
    filepath = os.path.join(DOCS_DIR, f"{case_data.get('fir_number')}_face_id.docx")
    doc.save(filepath)
    return filepath
