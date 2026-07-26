import pytest
import asyncio
import hashlib
from docx import Document
from app.core.redis import LazyRedisProxy
from app.db.connection import fetch_all, fetch_one
from app.services.document_gen import replace_in_paragraph
from app.services.prediction import RiskPredictor
from app.api.auth import create_access_token

@pytest.mark.asyncio
async def test_unauthenticated_endpoints_return_401(async_client):
    """Verify that unauthenticated routes enforce authentication and return 401."""
    # CCTV endpoints
    res_cameras = await async_client.get("/cctv/cameras")
    assert res_cameras.status_code == 401

    res_alert = await async_client.post("/cctv/alert", json={
        "camera_id": "CAM1",
        "source": "iccc",
        "alert_type": "loitering",
        "confidence": 0.9,
        "lat": 23.0,
        "lon": 72.0
    })
    assert res_alert.status_code == 401

    # Incidents SLA breaches endpoint
    res_sla = await async_client.get("/incident/sla_breaches")
    assert res_sla.status_code == 401

    # Patrol PCR endpoint
    res_pcr = await async_client.post("/patrol/pcr", json={
        "incident_type": "snatching",
        "location_text": "Station",
        "lat": 23.0,
        "lon": 72.0
    })
    assert res_pcr.status_code == 401

@pytest.mark.asyncio
async def test_redis_proxy_connection_reuse():
    """Verify LazyRedisProxy reuses client instance within the same event loop."""
    proxy = LazyRedisProxy()
    c1 = proxy._get_client()
    c2 = proxy._get_client()
    assert c1 is c2

@pytest.mark.asyncio
async def test_db_exception_reraising(db_session):
    """Verify fetch_all and fetch_one re-raise exceptions on invalid SQL queries."""
    with pytest.raises(Exception):
        await fetch_all(db_session, "SELECT * FROM non_existent_table_12345")

    with pytest.raises(Exception):
        await fetch_one(db_session, "SELECT * FROM non_existent_table_12345")

def test_document_gen_multi_run_formatting_preservation():
    """Verify tag replacement across multiple runs preserves run formatting."""
    doc = Document()
    para = doc.add_paragraph()
    r1 = para.add_run("Hello ")
    r2 = para.add_run("{{NAME")
    r2.bold = True
    r3 = para.add_run("}} world!")
    r3.italic = True

    ctx = {"NAME": "Alice"}
    replace_in_paragraph(para, ctx)

    assert para.text == "Hello Alice world!"
    assert para.runs[1].text == "Alice"
    assert para.runs[1].bold is True
    assert para.runs[2].text == " world!"
    assert para.runs[2].italic is True

@pytest.mark.asyncio
async def test_prediction_deterministic_hash(db_session):
    """Verify risk prediction ward hash is deterministic across multiple calls."""
    ward = "Ellisbridge"
    expected_hash = (int(hashlib.sha256(ward.encode('utf-8')).hexdigest(), 16) % 100) / 100.0
    assert 0.0 <= expected_hash < 1.0

    predictor = RiskPredictor.__new__(RiskPredictor)
    predictor._is_trained = True
    predictor.model = type("DummyModel", (), {"predict": lambda self, X: [50.0]})()
    r1 = await predictor.predict_zone_risk(ward, 14, 2, 7, db_session)
    r2 = await predictor.predict_zone_risk(ward, 14, 2, 7, db_session)
    assert r1 == r2
    assert r1 == 50.0 + (expected_hash * 10 - 5)
